"""
RAG Service Layer for Document Indexing and Retrieval.
This module handles all Haystack + ChromaDB operations.
"""

import os
import logging

from typing import List, Optional

from django.conf import settings

from haystack import Document as HaystackDocument
from haystack.components.embedders import SentenceTransformersDocumentEmbedder, SentenceTransformersTextEmbedder
from haystack.components.writers import DocumentWriter


from google import genai

from haystack_integrations.document_stores.chroma import ChromaDocumentStore
from haystack_integrations.components.retrievers.chroma import ChromaEmbeddingRetriever

from pypdf import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentParser:
    """Handles parsing of different document types."""

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from a PDF file."""
        try:
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """Read text from a TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    @classmethod
    def parse(cls, file_path: str, extension: str) -> str:
        """Parse document based on extension."""
        parsers = {
            'pdf': cls.parse_pdf,
            'docx': cls.parse_docx,
            'txt': cls.parse_txt,
        }
        parser = parsers.get(extension.lower())
        if not parser:
            raise ValueError(f"Unsupported file type: {extension}")
        return parser(file_path)


class TextChunker:
    """Simple text chunker for RAG."""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        if not text:
            return []

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = start + self.chunk_size

            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence-ending punctuation
                for punct in ['. ', '.\n', '! ', '? ']:
                    last_punct = text[start:end].rfind(punct)
                    if last_punct > self.chunk_size // 2:
                        end = start + last_punct + len(punct)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - self.chunk_overlap

        return chunks


class RagService:
    """
    Main RAG service class.
    Handles document indexing and query answering.
    """

    _instance = None
    _document_store = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance

    def __init__(self):
        if RagService._document_store is None:
            self._initialize_document_store()

    def _initialize_document_store(self):
        """Initialize ChromaDB document store."""
        persist_path = str(settings.CHROMA_PERSIST_DIRECTORY)
        os.makedirs(persist_path, exist_ok=True)

        RagService._document_store = ChromaDocumentStore(
            collection_name=settings.CHROMA_COLLECTION_NAME,
            persist_path=persist_path,
        )
        logger.info(f"Initialized ChromaDB at {persist_path}")

    @property
    def document_store(self):
        return RagService._document_store

    def index_document(self, document_model, user_id: int) -> bool:
        """
        Index a document into ChromaDB.
        
        Args:
            document_model: Django Document model instance
            user_id: ID of the owning user
            
        Returns:
            True if indexing succeeded, False otherwise
        """
        try:
            file_path = document_model.file.path
            extension = document_model.file_extension

            # Parse document
            text = DocumentParser.parse(file_path, extension)
            if not text.strip():
                logger.warning(f"No text extracted from {document_model.name}")
                return False

            # Chunk text
            chunker = TextChunker()
            chunks = chunker.chunk(text)

            # Create Haystack documents with metadata
            haystack_docs = []
            for i, chunk in enumerate(chunks):
                doc = HaystackDocument(
                    content=chunk,
                    meta={
                        'user_id': user_id,
                        'document_id': document_model.id,
                        'document_name': document_model.name,
                        'chunk_index': i,
                    }
                )
                haystack_docs.append(doc)

            # Create indexing pipeline
            embedder = SentenceTransformersDocumentEmbedder(
                model="sentence-transformers/all-MiniLM-L6-v2"
            )
            embedder.warm_up()

            writer = DocumentWriter(document_store=self.document_store)

            # Embed and write documents
            embedded_docs = embedder.run(documents=haystack_docs)
            writer.run(documents=embedded_docs['documents'])

            logger.info(f"Indexed {len(chunks)} chunks from {document_model.name}")
            return True

        except Exception as e:
            logger.error(f"Error indexing document {document_model.id}: {e}")
            return False

    def delete_document_embeddings(self, document_id: int, user_id: int):
        """
        Remove embeddings for a specific document.
        Called when a document is soft-deleted.
        """
        try:
            # ChromaDB filtering to find and delete documents
            # Note: This requires fetching IDs first, then deleting
            filters = {
                "meta.document_id": document_id,
                "meta.user_id": user_id
            }
            # Get document IDs to delete
            docs = self.document_store.filter_documents(filters=filters)
            if docs:
                doc_ids = [doc.id for doc in docs]
                self.document_store.delete_documents(document_ids=doc_ids)
                logger.info(f"Deleted {len(doc_ids)} embeddings for document {document_id}")
        except Exception as e:
            logger.error(f"Error deleting embeddings for document {document_id}: {e}")

    def answer_question(
        self,
        question: str,
        user_id: int,
        selected_document_ids: List[int],
        chat_history: Optional[List[dict]] = None
    ) -> str:
        """
        Answer a question using RAG.
        
        Args:
            question: The user's question
            user_id: ID of the user asking
            selected_document_ids: List of document IDs to search in
            chat_history: Optional list of previous messages for context
            
        Returns:
            The generated answer string
        """
        try:
            # Build filters for user and selected documents
            filters = {
                "operator": "AND",
                "conditions": [
                    {"field": "meta.user_id", "operator": "==", "value": user_id},
                    {"field": "meta.document_id", "operator": "in", "value": selected_document_ids},
                ]
            }

            # Create text embedder for query
            text_embedder = SentenceTransformersTextEmbedder(
                model="sentence-transformers/all-MiniLM-L6-v2"
            )
            text_embedder.warm_up()

            # Embed the query
            query_embedding_result = text_embedder.run(text=question)
            query_embedding = query_embedding_result['embedding']

            # Retrieve relevant documents
            retriever = ChromaEmbeddingRetriever(document_store=self.document_store)
            retrieval_result = retriever.run(
                query_embedding=query_embedding,
                filters=filters,
                top_k=5
            )
            docs = retrieval_result['documents']

            if not docs:
                return "I couldn't find any relevant information in the selected documents."

            # Build context from retrieved documents
            context_parts = []
            for doc in docs:
                context_parts.append(doc.content)
            context = "\n\n---\n\n".join(context_parts)

            # Build chat history context
            history_text = ""
            if chat_history:
                history_parts = []
                for msg in chat_history[-10:]:  # Last 10 messages
                    role = msg.get('role', 'user').capitalize()
                    content = msg.get('content', '')
                    history_parts.append(f"{role}: {content}")
                history_text = "\n".join(history_parts)

            # Create the prompt
            prompt_template = """You are a helpful assistant answering questions based on the provided documents.

Recent conversation:
{history}

Relevant document excerpts:
{context}

User question: {question}

Please provide a clear, accurate answer based on the document excerpts above. If the documents don't contain enough information to answer the question, say so clearly.

Answer:"""

            # Fill the template manually as we're not using PromptBuilder anymore for the final call
            prompt = prompt_template.format(
                history=history_text if history_text else "No previous history.",
                context=context,
                question=question
            )

            # Generate answer using official google-genai library
            api_key = settings.GOOGLE_API_KEY
            if not api_key or api_key == 'your-google-api-key-here':
                return "Error: Google API key is not configured. Please set GOOGLE_API_KEY in your .env file."

            # Initialize official GenAI client
            client = genai.Client(api_key=api_key)
            
            # Use gemini-2.5-flash as requested
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt
            )

            if response and response.text:
                return response.text
            else:
                return "I couldn't generate a response. Please try again."

        except Exception as e:
            logger.error(f"Error answering question: {e}")
            return f"An error occurred while processing your question: {str(e)}"


# Singleton instance
rag_service = RagService()
